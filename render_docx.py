#!/usr/bin/env python3
"""
render_docx.py — Convert a markdown report to a .docx file.

Minimal markdown → docx converter supporting:
- Headings (# through ####)
- Paragraphs with inline **bold**, *italic*, `code`
- Bullet lists (- / *)
- Numbered lists (1. 2. ...)
- Code blocks (``` fenced)
- Tables (GitHub-style pipe syntax)
- Horizontal rules (---)
- Checkboxes ([ ], [x])
- Blockquotes (>)

Usage:
    python render_docx.py <input.md> <output.docx>
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.+)$")
NUMBERED_RE = re.compile(r"^(\s*)\d+\.\s+(.+)$")
CHECKBOX_RE = re.compile(r"^\s*-\s+\[( |x|X)\]\s+(.+)$")
INLINE_TOKEN_RE = re.compile(r"(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+\]\([^)]+\))")
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def add_inline(paragraph, text: str) -> None:
    """Render inline markdown (bold, italic, code, links) into a paragraph as runs."""
    # Strip image syntax (we don't embed images in this simple renderer)
    text = IMAGE_RE.sub(r"[\1]", text)
    pieces = INLINE_TOKEN_RE.split(text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
        elif piece.startswith("*") and piece.endswith("*") and not piece.startswith("**"):
            run = paragraph.add_run(piece[1:-1])
            run.italic = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif piece.startswith("[") and "](" in piece and piece.endswith(")"):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", piece)
            if m:
                add_hyperlink(paragraph, m.group(1), m.group(2))
        else:
            paragraph.add_run(piece)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def render_table(doc: Document, rows: list[str]) -> None:
    """Render a GitHub-style pipe table. rows includes header + separator + data rows."""
    if not rows:
        return
    parsed_rows: list[list[str]] = []
    for line in rows:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        parsed_rows.append(cells)
    # Drop separator row(s) like |---|---|
    data = [r for r in parsed_rows if not all(re.match(r"^:?-+:?$", c) for c in r)]
    if not data:
        return
    col_count = max(len(r) for r in data)
    table = doc.add_table(rows=len(data), cols=col_count)
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(data):
        for ci in range(col_count):
            cell_text = row[ci] if ci < len(row) else ""
            cell = table.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            if ri == 0:
                run = para.add_run(cell_text)
                run.bold = True
            else:
                add_inline(para, cell_text)
    doc.add_paragraph()


def convert(md_path: Path, docx_path: Path) -> None:
    md = md_path.read_text(encoding="utf-8")
    lines = md.splitlines()

    doc = Document()

    # Base style tuning
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)

    # Header styles — a touch of visual hierarchy
    for lvl, size in [(1, 20), (2, 16), (3, 13), (4, 12)]:
        try:
            style = doc.styles[f"Heading {lvl}"]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
        except KeyError:
            pass

    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code fence
        if stripped.startswith("```"):
            if in_code:
                # Flush code block
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                # Grey box-ish background (simple tint via shade)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F3F3F3")
                para.paragraph_format.left_indent = Inches(0.2)
                try:
                    para._p.get_or_add_pPr().append(shd)
                except Exception:
                    pass
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule
        if stripped in {"---", "***", "___"}:
            p = doc.add_paragraph()
            p.add_run("─" * 40)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Heading
        m = HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            heading = doc.add_heading("", level=level)
            add_inline(heading, m.group(2))
            i += 1
            continue

        # Table
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(lines[i])
                i += 1
            render_table(doc, table_rows)
            continue

        # Checkbox list
        m = CHECKBOX_RE.match(line)
        if m:
            checked = m.group(1).lower() == "x"
            mark = "☑" if checked else "☐"
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"{mark} ")
            add_inline(para, m.group(2))
            i += 1
            continue

        # Bullet
        m = BULLET_RE.match(line)
        if m:
            para = doc.add_paragraph(style="List Bullet")
            add_inline(para, m.group(2))
            i += 1
            continue

        # Numbered
        m = NUMBERED_RE.match(line)
        if m:
            para = doc.add_paragraph(style="List Number")
            add_inline(para, m.group(2))
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            para = doc.add_paragraph()
            run = para.add_run(stripped.lstrip("> ").strip())
            run.italic = True
            para.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Paragraph (collect continuation lines until blank/structural)
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip():
                break
            if HEADING_RE.match(nxt) or nxt.strip().startswith("```") \
                or BULLET_RE.match(nxt) or NUMBERED_RE.match(nxt) \
                or nxt.strip().startswith("|") or nxt.strip().startswith(">") \
                or nxt.strip() in {"---", "***", "___"}:
                break
            para_lines.append(nxt)
            i += 1
        para = doc.add_paragraph()
        add_inline(para, " ".join(l.strip() for l in para_lines))

    doc.save(str(docx_path))


def main() -> None:
    if len(sys.argv) != 3:
        print("Usage: python render_docx.py <input.md> <output.docx>", file=sys.stderr)
        sys.exit(1)
    md_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])
    if not md_path.exists():
        print(f"Input not found: {md_path}", file=sys.stderr)
        sys.exit(1)
    convert(md_path, docx_path)
    print(f"Wrote {docx_path}", file=sys.stderr)


if __name__ == "__main__":
    main()
